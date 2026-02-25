# main.py
import io
import uuid
import re
import hashlib
import docx
import os
import json
import faiss
import threading
from datetime import datetime
from typing import List, Optional

from fastapi import FastAPI, UploadFile, File, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
import pdfplumber
import pytesseract
from PIL import Image
from pypdf import PdfReader

# Importing your custom logic
from vector_store import search_duplicate, add_to_index
from fraud_detection import detect_pii, analyze_metadata, extract_advanced_entities
from image_forensics import detect_tampering, get_image_phash
from pydantic import BaseModel, Field

# ------------------------------------------------------------------
# SYSTEM RESET LOGIC (MANUAL ONLY – SAFE FOR CLOUD)
# ------------------------------------------------------------------

def reset_system_data():
    global db
    db.clear()

    try:
        with open("hash.json", "w") as f:
            json.dump({}, f)
    except Exception:
        pass

    try:
        with open("sha256.json", "w") as f:
            json.dump({}, f)
    except Exception:
        pass

    try:
        dimension = 384
        index = faiss.IndexFlatIP(dimension)
        faiss.write_index(index, "docs.index")
    except Exception:
        pass

    return True

# ------------------------------------------------------------------
# APPLICATION INITIALIZATION
# ------------------------------------------------------------------

app = FastAPI(
    title="AP FraudShield Final",
    version="1.0.0",
)

# ------------------------------------------------------------------
# CORS (FIXED – USE ENV VARIABLE PROPERLY)
# ------------------------------------------------------------------

# Default origins include both production and local development
DEFAULT_ORIGINS = ",".join([
    "https://finance-ai-project-eight.vercel.app",
    "http://localhost:3000",
    "http://127.0.0.1:3000",
])

ALLOWED_ORIGINS = os.getenv("CORS_ORIGINS", DEFAULT_ORIGINS).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    
)

# In-memory DB with thread-safe access
db = {}
db_lock = threading.Lock()

# Maximum file size (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}

# ------------------------------------------------------------------
# SCHEMAS (UNCHANGED)
# ------------------------------------------------------------------

class AnomalyItem(BaseModel):
    type: str
    description: str
    confidence: float = Field(..., ge=0, le=1)

class ScanResult(BaseModel):
    file_id: str
    filename: str
    file_url: str
    fraud_score: int
    severity: str
    anomalies: List[AnomalyItem]
    text_content: str
    extracted_tables: List
    processing_time: int
    confidence: float = Field(..., ge=0, le=1)

class UploadResponse(BaseModel):
    task_id: str
    message: str

class AlertRequest(BaseModel):
    message: str

class AlertResponse(BaseModel):
    status: str

# ------------------------------------------------------------------
# HELPERS (LOGIC UNCHANGED – ONLY PDF OCR REMOVED)
# ------------------------------------------------------------------

def clean_text(text: str) -> str:
    cleaned = re.sub(r"[^\w\s.,:/-]", " ", text)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()

def extract_text_from_file(content: bytes, filename: str) -> str:
    name = filename.lower()

    if name.endswith((".jpg", ".jpeg", ".png")):
        try:
            image = Image.open(io.BytesIO(content)).convert("L")
            return clean_text(pytesseract.image_to_string(image))
        except Exception:
            return ""

    if name.endswith(".docx") or name.endswith(".doc"):
        try:
            doc = docx.Document(io.BytesIO(content))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            return clean_text(" ".join(full_text))
        except Exception:
            return ""

    if name.endswith(".pdf"):
        # Primary: extract selectable text layer
        extracted = ""
        try:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                if i >= 5:
                    break
                pages_text.append(page.extract_text() or "")
            extracted = clean_text(" ".join(pages_text))
        except Exception:
            extracted = ""

        # OCR fallback: if text layer is too sparse, run Tesseract on first page image
        if len(extracted.strip()) < 50:
            try:
                from pdf2image import convert_from_bytes
                pages = convert_from_bytes(content, dpi=200, first_page=1, last_page=1)
                if pages:
                    ocr_text = pytesseract.image_to_string(pages[0].convert("L"))
                    extracted = clean_text(ocr_text)
            except Exception:
                pass

        return extracted

    return ""

# ------------------------------------------------------------------
# ADMIN RESET ROUTE (UNCHANGED)
# ------------------------------------------------------------------

@app.get("/api/v1/admin/reset")
async def manual_reset(key: str):
    """Admin reset endpoint - requires secret key from environment variable."""
    admin_key = os.getenv("ADMIN_RESET_KEY", "ap_finance_2025")
    if key != admin_key:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Unauthorized")
    reset_system_data()
    return {"status": "success", "message": "Backend data wiped."}

# ------------------------------------------------------------------
# DASHBOARD (FAST – SAFE FOR CLOUD)
# ------------------------------------------------------------------

@app.get("/api/v1/dashboard/stats")
def get_dashboard_stats():
    return {
        "summary": {
            "total_scanned": 14205,
            "fraud_detected": 45,
            "savings_in_crores": 1.2,
        },
        "weekly_activity": [
            {"day": "Mon", "uploads": 120, "fraud": 2},
            {"day": "Tue", "uploads": 150, "fraud": 5},
            {"day": "Wed", "uploads": 180, "fraud": 1},
            {"day": "Thu", "uploads": 90, "fraud": 0},
            {"day": "Fri", "uploads": 200, "fraud": 8},
            {"day": "Sat", "uploads": 50, "fraud": 0},
            {"day": "Sun", "uploads": 30, "fraud": 0},
        ],
        "recent_scans": [
            {"id": "1", "filename": "invoice_992.pdf", "status": "safe", "timestamp": "2 mins ago"},
            {"id": "2", "filename": "contract_v2.docx", "status": "warning", "timestamp": "5 mins ago"},
        ],
    }

# ------------------------------------------------------------------
# MAIN SCAN ROUTE (UNCHANGED LOGIC)
# ------------------------------------------------------------------

@app.post("/api/v1/scan/upload")
async def upload_scan(file: UploadFile = File(...)):
    """Upload and scan a document for fraud detection."""
    start_time = datetime.now()
    
    # Validate filename exists
    filename = file.filename
    if not filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is required"
        )
    
    # Validate file extension
    file_ext = os.path.splitext(filename.lower())[1]
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type '{file_ext}' not supported. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Read and validate file content
    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Empty file received"
        )
    
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE // (1024 * 1024)}MB"
        )
    
    task_id = str(uuid.uuid4())

    # Compute SHA-256 fingerprint of raw file bytes (Layer 0 duplicate check)
    sha256_hash = hashlib.sha256(content).hexdigest()

    text = extract_text_from_file(content, filename)
    entities = extract_advanced_entities(text)

    tables = []
    if filename.lower().endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                tables = [p.extract_table() for p in pdf.pages if p.extract_table()]
        except Exception:
            pass

    img_hash = get_image_phash(content, filename)
    is_dup, dup_score = search_duplicate(text, img_hash, sha256_hash)
    tamper_msg, tamper_conf = detect_tampering(content, filename)
    meta_issue, meta_conf = analyze_metadata(content, text)
    pii_found, pii_conf = detect_pii(text)

    fraud_score = 0
    anomalies = []

    if tamper_msg:
        fraud_score = max(fraud_score, 90)
        anomalies.append({"type": "Forensic Tampering", "description": tamper_msg, "confidence": tamper_conf})

    if meta_issue:
        fraud_score = max(fraud_score, 85)
        anomalies.append({"type": "Metadata Fraud", "description": meta_issue, "confidence": meta_conf})

    if is_dup:
        fraud_score = 100
        anomalies.append({"type": "Duplicate Discovery", "description": "Visual or text match found.", "confidence": dup_score})

    if pii_found:
        anomalies.append({"type": "PII Detected", "description": f"Contains: {pii_found}", "confidence": pii_conf})
        if fraud_score < 30:
            fraud_score += 20

    severity = "CRITICAL" if fraud_score >= 70 else "WARNING" if fraud_score >= 30 else "SAFE"

    if not is_dup:
        add_to_index(text, img_hash, sha256_hash)

    overall_confidence = sum(a["confidence"] for a in anomalies) / len(anomalies) if anomalies else 0.0

    result = {
        "file_id": task_id,
        "filename": filename,
        "file_url": f"/api/v1/files/{task_id}",  # Virtual file reference
        "fraud_score": min(100, fraud_score),
        "severity": severity,
        # Duplicate detection result — consumed by the frontend AnalysisPanel
        "is_duplicate": is_dup,
        "duplicate_source_id": None,  # Would require a persistent DB to track original doc ID
        "anomalies": anomalies,
        "text_content": text,
        "entities": entities,
        "extracted_tables": tables,
        "processing_time": int((datetime.now() - start_time).total_seconds() * 1000),
        "confidence": round(overall_confidence, 4),
        "status": "completed",
        "scanned_at": datetime.now().isoformat(),
    }

    with db_lock:
        db[task_id] = result
    return {"task_id": task_id, "message": "Unified fraud analysis concluded."}

# ------------------------------------------------------------------
# RESULT + HEALTH
# ------------------------------------------------------------------

@app.get("/api/v1/scan/result/{task_id}")
async def get_result(task_id: str):
    """Retrieve scan result by task ID."""
    with db_lock:
        result = db.get(task_id)
    
    if result is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Scan result with task_id '{task_id}' not found"
        )
    return result

@app.get("/health")
@app.get("/api/v1/health")
def health_check():
    return {"status": "healthy", "version": "1.0.0"}

@app.post("/api/v1/admin/trigger-alert", response_model=AlertResponse)
def trigger_alert(payload: AlertRequest):
    return {"status": "sent"}
